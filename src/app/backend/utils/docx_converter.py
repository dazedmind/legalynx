"""
DOCX to PDF Converter Utility

This module provides functionality to convert DOCX files to PDF format
using python-docx for reading and reportlab for PDF generation.
"""

import os
import re
from typing import List, Optional
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT


class DocxToPdfConverter:
    """
    A utility class for converting DOCX files to PDF format.
    """
    
    def __init__(self):
        """Initialize the converter with default styles."""
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()
    
    def _setup_custom_styles(self):
        """Setup custom paragraph styles for better formatting."""
        self.normal_style = ParagraphStyle(
            'CustomNormal',
            parent=self.styles['Normal'],
            fontSize=11,
            spaceAfter=12,
            alignment=TA_LEFT,
            fontName='Helvetica'
        )
        
        self.heading_style = ParagraphStyle(
            'CustomHeading',
            parent=self.styles['Heading1'],
            fontSize=14,
            spaceAfter=12,
            spaceBefore=12,
            alignment=TA_LEFT,
            fontName='Helvetica-Bold'
        )
    
    def _clean_text(self, text: str) -> str:
        """
        Clean text for PDF generation by handling special characters.
        
        Args:
            text: Raw text from DOCX
            
        Returns:
            Cleaned text safe for PDF generation
        """
        if not text:
            return ""
        
        # Replace problematic Unicode characters
        replacements = {
            '\u2019': "'",      # Right single quotation mark
            '\u2018': "'",      # Left single quotation mark
            '\u201c': '"',      # Left double quotation mark
            '\u201d': '"',      # Right double quotation mark
            '\u2013': '-',      # En dash
            '\u2014': '-',      # Em dash
            '\u2026': '...',    # Horizontal ellipsis
            '\u00a0': ' ',      # Non-breaking space
            '\u2022': '•',      # Bullet point
        }
        
        for unicode_char, replacement in replacements.items():
            text = text.replace(unicode_char, replacement)
        
        # Remove any remaining problematic characters
        text = re.sub(r'[^\x00-\x7F]+', ' ', text)
        
        # Escape XML characters for reportlab
        text = text.replace('&', '&amp;')
        text = text.replace('<', '&lt;')
        text = text.replace('>', '&gt;')
        
        return text.strip()
    
    def _is_heading(self, paragraph, text: str) -> bool:
        """
        Determine if a paragraph should be treated as a heading.
        
        Args:
            paragraph: DOCX paragraph object
            text: Paragraph text
            
        Returns:
            True if paragraph should be treated as heading
        """
        if not text:
            return False
        
        # Check style name
        if hasattr(paragraph, 'style') and paragraph.style.name.startswith('Heading'):
            return True
        
        # Simple heuristics for heading detection
        return (
            len(text) < 100 and 
            (text.isupper() or 
             text.endswith(':') or
             text.count(' ') < 8)  # Short lines are likely headings
        )
    
    def _process_paragraphs(self, doc: DocxDocument) -> List:
        """
        Process all paragraphs from DOCX document.
        
        Args:
            doc: DOCX document object
            
        Returns:
            List of reportlab story elements
        """
        story = []
        
        for paragraph in doc.paragraphs:
            text = self._clean_text(paragraph.text)
            
            if not text:
                # Add small space for empty paragraphs
                story.append(Spacer(1, 6))
                continue
            
            try:
                # Determine paragraph style
                if self._is_heading(paragraph, text):
                    para = Paragraph(text, self.heading_style)
                else:
                    para = Paragraph(text, self.normal_style)
                
                story.append(para)
                
            except Exception as e:
                print(f"⚠️ Warning: Could not format paragraph, using fallback: {str(e)[:50]}")
                try:
                    # Fallback: ASCII-only text
                    ascii_text = ''.join(c for c in text if ord(c) < 128)
                    if ascii_text.strip():
                        para = Paragraph(ascii_text, self.normal_style)
                        story.append(para)
                except Exception:
                    print(f"⚠️ Skipping problematic paragraph")
                    continue
        
        return story
    
    def _process_tables(self, doc: DocxDocument) -> List:
        """
        Process all tables from DOCX document.
        
        Args:
            doc: DOCX document object
            
        Returns:
            List of reportlab story elements for tables
        """
        story = []
        
        for table in doc.tables:
            # Add spacer before table
            story.append(Spacer(1, 12))
            
            for row in table.rows:
                # Extract cell text and join with separator
                cell_texts = []
                for cell in row.cells:
                    cell_text = self._clean_text(cell.text)
                    if cell_text:
                        cell_texts.append(cell_text)
                
                if cell_texts:
                    row_text = " | ".join(cell_texts)
                    try:
                        para = Paragraph(row_text, self.normal_style)
                        story.append(para)
                    except Exception:
                        print("⚠️ Skipping problematic table row")
                        continue
            
            # Add spacer after table
            story.append(Spacer(1, 12))
        
        return story
    
    def convert(self, docx_path: str, pdf_path: str) -> bool:
        """
        Convert DOCX file to PDF.
        
        Args:
            docx_path: Path to input DOCX file
            pdf_path: Path where PDF should be saved
            
        Returns:
            True if conversion successful, False otherwise
            
        Raises:
            Exception: If conversion fails completely
        """
        try:
            print(f"🔄 Converting DOCX to PDF: {os.path.basename(docx_path)}")
            
            # Validate input file
            if not os.path.exists(docx_path):
                raise FileNotFoundError(f"DOCX file not found: {docx_path}")
            
            # Read DOCX document
            doc = DocxDocument(docx_path)
            
            # Create PDF document
            pdf_doc = SimpleDocTemplate(
                pdf_path,
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18
            )
            
            # Process content
            story = []
            
            # Add paragraphs
            paragraph_content = self._process_paragraphs(doc)
            story.extend(paragraph_content)
            
            # Add tables
            table_content = self._process_tables(doc)
            story.extend(table_content)
            
            # Fallback content if nothing was extracted
            if not story:
                story.append(Paragraph(
                    "Document converted from DOCX format", 
                    self.normal_style
                ))
                story.append(Paragraph(
                    "The original document may have been empty or contained unsupported content.", 
                    self.normal_style
                ))
            
            # Generate PDF
            pdf_doc.build(story)
            
            # Verify PDF was created
            if not os.path.exists(pdf_path):
                raise Exception("PDF file was not created")
            
            print(f"✅ Successfully converted to PDF: {os.path.basename(pdf_path)}")
            return True
            
        except Exception as e:
            print(f"❌ DOCX conversion failed: {str(e)}")
            # Clean up partial PDF file if it exists
            if os.path.exists(pdf_path):
                try:
                    os.remove(pdf_path)
                except:
                    pass
            raise Exception(f"DOCX conversion failed: {str(e)}")


class SimpleDocxConverter:
    """
    A simplified DOCX to PDF converter as a fallback option.
    """
    
    @staticmethod
    def convert(docx_path: str, pdf_path: str) -> bool:
        """
        Simple conversion with minimal formatting.
        
        Args:
            docx_path: Path to input DOCX file
            pdf_path: Path where PDF should be saved
            
        Returns:
            True if conversion successful
            
        Raises:
            Exception: If conversion fails
        """
        try:
            print(f"🔄 Simple DOCX conversion: {os.path.basename(docx_path)}")
            
            # Read DOCX
            doc = DocxDocument(docx_path)
            
            # Extract all text
            all_text = []
            
            # Get paragraph text
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    # Clean for ASCII compatibility
                    clean_text = ''.join(c if ord(c) < 128 else ' ' for c in text)
                    clean_text = clean_text.replace('&', 'and').replace('<', '[').replace('>', ']')
                    if clean_text.strip():
                        all_text.append(clean_text.strip())
            
            # Get table text
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        clean_row = ''.join(c if ord(c) < 128 else ' ' for c in row_text)
                        all_text.append(clean_row)
            
            # Create simple PDF
            pdf_doc = SimpleDocTemplate(pdf_path, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            if not all_text:
                all_text = ["Converted from DOCX file", "Content may be empty or unsupported."]
            
            for text in all_text:
                if text.strip():
                    para = Paragraph(text, styles['Normal'])
                    story.append(para)
                    story.append(Spacer(1, 12))
            
            pdf_doc.build(story)
            print(f"✅ Simple conversion completed: {os.path.basename(pdf_path)}")
            return True
            
        except Exception as e:
            print(f"❌ Simple conversion failed: {str(e)}")
            if os.path.exists(pdf_path):
                try:
                    os.remove(pdf_path)
                except:
                    pass
            raise Exception(f"Simple DOCX conversion failed: {str(e)}")


# Convenience functions
def convert_docx_to_pdf(docx_path: str, pdf_path: str, use_simple: bool = False) -> bool:
    """
    Convert DOCX to PDF using the appropriate converter.
    
    Args:
        docx_path: Path to input DOCX file
        pdf_path: Path where PDF should be saved
        use_simple: If True, use simple converter; if False, try advanced first
        
    Returns:
        True if conversion successful
        
    Raises:
        Exception: If conversion fails completely
    """
    if use_simple:
        return SimpleDocxConverter.convert(docx_path, pdf_path)
    
    try:
        # Try advanced converter first
        converter = DocxToPdfConverter()
        return converter.convert(docx_path, pdf_path)
    except Exception as e:
        print(f"⚠️ Advanced converter failed, trying simple converter: {str(e)}")
        try:
            return SimpleDocxConverter.convert(docx_path, pdf_path)
        except Exception as simple_error:
            print(f"❌ Both converters failed")
            raise Exception(f"DOCX conversion failed. Advanced: {str(e)}, Simple: {str(simple_error)}")


def is_docx_file(filename: str) -> bool:
    """
    Check if a file is a DOCX file based on extension.
    
    Args:
        filename: Name of the file to check
        
    Returns:
        True if file appears to be DOCX format
    """
    if not filename:
        return False
    
    return filename.lower().endswith(('.docx', '.doc'))


def validate_docx_file(docx_path: str) -> dict:
    """
    Validate a DOCX file and return information about it.
    
    Args:
        docx_path: Path to DOCX file
        
    Returns:
        Dictionary with validation results
    """
    result = {
        'is_valid': False,
        'error': None,
        'paragraph_count': 0,
        'table_count': 0,
        'has_content': False
    }
    
    try:
        if not os.path.exists(docx_path):
            result['error'] = 'File does not exist'
            return result
        
        doc = DocxDocument(docx_path)
        
        # Count content
        result['paragraph_count'] = len(doc.paragraphs)
        result['table_count'] = len(doc.tables)
        
        # Check for actual content
        has_text = any(p.text.strip() for p in doc.paragraphs)
        has_table_content = any(
            any(cell.text.strip() for cell in row.cells)
            for table in doc.tables
            for row in table.rows
        )
        
        result['has_content'] = has_text or has_table_content
        result['is_valid'] = True
        
    except Exception as e:
        result['error'] = str(e)
    
    return result