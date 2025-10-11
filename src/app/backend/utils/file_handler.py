import os
import fitz
import shutil
from typing import List, Optional, Dict, Any
from llama_index.core import Document
from datetime import datetime
import re
import os

# DOCX text extraction
try:
    from docx import Document as DocxDocument
    DOCX_TEXT_EXTRACTION = True
    print("✅ DOCX text extraction available (python-docx)")
except ImportError:
    DOCX_TEXT_EXTRACTION = False
    print("⚠️ DOCX text extraction not available - install python-docx")

def extract_text_from_docx(docx_path: str) -> List[Document]:
    """
    Extract text directly from DOCX without conversion to PDF.
    This is the most reliable method when PDF conversion fails.
    """
    if not DOCX_TEXT_EXTRACTION:
        raise Exception("python-docx not installed. Install with: pip install python-docx")
    
    print("📄 Extracting text directly from DOCX...")
    documents = []
    
    try:
        # Load DOCX document
        doc = DocxDocument(docx_path)
        
        # Extract text from all paragraphs
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text.strip())
        
        # Extract text from tables
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    table_text.append(row_text)
        
        # Combine all text
        all_content = full_text + table_text
        combined_text = "\n".join(all_content)
        
        if not combined_text.strip():
            combined_text = "No extractable text found in DOCX document."
        
        # Create a single document (DOCX doesn't have pages like PDF)
        documents.append(Document(
            text=combined_text,
            metadata={
                "file_name": os.path.basename(docx_path),
                "page_number": 1,
                "total_pages": 1,
                "extraction_method": "docx_direct",
                "original_format": "docx",
                "paragraph_count": len(full_text),
                "table_count": len([t for t in doc.tables])
            }
        ))
        
        print(f"✅ Successfully extracted text from DOCX")
        print(f"📊 Total characters extracted: {len(combined_text)}")
        print(f"📄 Paragraphs: {len(full_text)}, Tables: {len([t for t in doc.tables])}")
        
        return documents
        
    except Exception as e:
        print(f"❌ Error extracting text from DOCX: {e}")
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")

def get_next_sequential_number(upload_dir: str, title: Optional[str] = None, client_name: Optional[str] = None) -> int:
    """
    Get the next sequential number for file naming.
    """
    if not title or not client_name:
        return 1
    
    # Create the pattern to search for
    pattern = f"{title}_{client_name}_"
    
    # Ensure directory exists
    os.makedirs(upload_dir, exist_ok=True)
    
    # Find existing files with this pattern
    existing_numbers = []
    for filename in os.listdir(upload_dir):
        if filename.startswith(pattern):
            # Extract the number part
            try:
                # Get the part after the last underscore and before the extension
                name_without_ext = os.path.splitext(filename)[0]
                number_part = name_without_ext.split('_')[-1]
                if number_part.isdigit():
                    existing_numbers.append(int(number_part))
            except (ValueError, IndexError):
                continue
    
    # Return the next number
    return max(existing_numbers) + 1 if existing_numbers else 1

def generate_filename(
    original_filename: str,
    naming_option: str,
    title: Optional[str] = None,
    client_name: Optional[str] = None,
    counter: Optional[int] = None
) -> str:
    """
    Generate filename based on naming option and user settings.
    For add_client_name: surname at the end in format YYYYMMDD_DOCUMENTTYPE_SURNAME.ext
    """
    file_ext = os.path.splitext(original_filename)[1].lower()
    
    def clean_for_filename(text: str) -> str:
        """Clean text to be safe for filenames - uppercase with dashes"""
        if not text:
            return ""
        
        cleaned = re.sub(r'[^\w\s-]', '', text)
        cleaned = re.sub(r'\s+', ' ', cleaned).strip()
        
        words_to_remove = ['page', 'unknown', 'document', 'file']
        words = cleaned.split()
        filtered_words = [w for w in words if w.lower() not in words_to_remove]
        
        if filtered_words:
            return '-'.join(filtered_words).upper()
        else:
            return re.sub(r'\s+', '-', cleaned).upper()
    
    def extract_surname(text: str) -> str:
        """Extract only the surname (last word) from client name."""
        if not text:
            return ""
        words = text.strip().split()
        return words[-1].upper() if words else ""
    
    if naming_option == "keep_original":
        return original_filename
    
    elif naming_option == "add_timestamp":
        if not title:
            print("⚠️ Missing title for timestamp naming, using original")
            return original_filename
        
        date_str = datetime.now().strftime("%Y%m%d")
        clean_title = clean_for_filename(title)
        return f"{date_str}_{clean_title}{file_ext}"
    
    elif naming_option == "add_client_name":
        if not title or not client_name:
            print("⚠️ Missing title or client_name for client naming, using original")
            return original_filename
        
        # Format: YYYYMMDD_DOCUMENTTYPE_SURNAME.ext (surname at the END)
        date_str = datetime.now().strftime("%Y%m%d")
        clean_title = clean_for_filename(title)
        surname = extract_surname(client_name)
        return f"{date_str}_{clean_title}_{surname}{file_ext}"
    
    else:
        print(f"⚠️ Unknown naming option: {naming_option}, using original")
        return original_filename


def generate_fallback_filename(
    original_filename: str,
    naming_option: str,
    user_title: Optional[str] = None,
    user_client_name: Optional[str] = None,
    counter: Optional[int] = None
) -> str:
    """Enhanced fallback filename generation - surname at the end for add_client_name."""
    file_ext = os.path.splitext(original_filename)[1].lower()
    
    def clean_for_filename(text: str) -> str:
        if not text:
            return ""
        
        cleaned = re.sub(r'[^\w\s-]', '', text)
        cleaned = re.sub(r'\s+', ' ', cleaned).strip()
        
        words_to_remove = ['page', 'unknown', 'document', 'file']
        words = cleaned.split()
        filtered_words = [w for w in words if w.lower() not in words_to_remove]
        
        if filtered_words:
            return '-'.join(filtered_words).upper()[:30]
        else:
            return re.sub(r'\s+', '-', cleaned).upper()[:30]
    
    def extract_surname(text: str) -> str:
        """Extract only the surname (last word) from client name."""
        if not text:
            return ""
        words = text.strip().split()
        return words[-1].upper() if words else ""
    
    if naming_option == "add_timestamp" and user_title:
        date_str = datetime.now().strftime("%Y%m%d")
        clean_title = clean_for_filename(user_title)
        return f"{date_str}_{clean_title}{file_ext}"
    
    elif naming_option == "add_client_name" and user_title and user_client_name:
        date_str = datetime.now().strftime("%Y%m%d")
        clean_title = clean_for_filename(user_title)
        surname = extract_surname(user_client_name)
        # Format: YYYYMMDD_DOCUMENTTYPE_SURNAME.ext
        return f"{date_str}_{clean_title}_{surname}{file_ext}"
    
    elif naming_option in ["add_timestamp", "add_client_name"]:
        extracted_name = extract_name_from_filename(original_filename)
        extracted_type = extract_type_from_filename(original_filename)
        
        if naming_option == "add_timestamp" and extracted_type:
            date_str = datetime.now().strftime("%Y%m%d")
            return f"{date_str}_{clean_for_filename(extracted_type)}{file_ext}"
        
        elif naming_option == "add_client_name" and (extracted_name or extracted_type):
            date_str = datetime.now().strftime("%Y%m%d")
            parts = [date_str]
            
            if extracted_type:
                parts.append(clean_for_filename(extracted_type))
            
            if extracted_name:
                surname = extract_surname(extracted_name)
                parts.append(surname)
            
            if len(parts) >= 2:
                return f"{'_'.join(parts)}{file_ext}"
    
    print(f"⚠️ No valid naming configuration, keeping original: {original_filename}")
    return original_filename

def save_uploaded_file(
    file_content: bytes, 
    filename: str, 
    naming_option: str,  # REQUIRED - no default value
    upload_dir: str = "sample_docs",
    rag_system: Optional[Dict[str, Any]] = None,
    user_title: Optional[str] = None,
    user_client_name: Optional[str] = None,
    counter: Optional[int] = None
) -> str:
    """
    Save uploaded file with intelligent naming using RAG analysis.
    
    Args:
        file_content: The file content as bytes
        filename: Original filename
        naming_option: How to name the file (REQUIRED - from user settings)
        upload_dir: Directory to save the file
        rag_system: RAG system for intelligent naming (optional)
        user_title: User-provided title (fallback)
        user_client_name: User-provided client name (fallback)
        counter: Sequential counter for numbering
    
    Returns:
        str: Path to the saved file
    """
    print(f"🔧 File naming function called with:")
    print(f"   - naming_option: {naming_option}")
    print(f"   - user_title: {user_title}")
    print(f"   - user_client_name: {user_client_name}")
    print(f"   - rag_system available: {rag_system is not None}")
    
    # Create directory if it doesn't exist
    os.makedirs(upload_dir, exist_ok=True)
    
    # If no RAG system provided or keep_original, use simple naming
    if not rag_system or naming_option == "keep_original":
        print(f"📁 Using simple file naming (naming_option: {naming_option})...")
        
        # Generate filename using existing logic
        new_filename = generate_filename(
            original_filename=filename,
            naming_option=naming_option,
            title=user_title,
            client_name=user_client_name,
            counter=counter
        )
        
        # Create full file path
        file_path = os.path.join(upload_dir, new_filename)
        
        # Handle file conflicts
        if os.path.exists(file_path) and naming_option == "keep_original":
            base_name, ext = os.path.splitext(new_filename)
            conflict_counter = 1
            while os.path.exists(file_path):
                file_path = os.path.join(upload_dir, f"{base_name}_{conflict_counter}{ext}")
                conflict_counter += 1
            new_filename = os.path.basename(file_path)
            print(f"   Conflict resolved: {new_filename}")
        
        # Save the file
        try:
            with open(file_path, 'wb') as f:
                f.write(file_content)
            
            print(f"✅ File saved with simple naming:")
            print(f"   Original: {filename}")
            print(f"   Final: {new_filename}")
            print(f"   Path: {file_path}")
            print(f"   Size: {len(file_content)} bytes")
            
            return file_path
            
        except Exception as e:
            print(f"❌ Failed to save file: {e}")
            raise Exception(f"Failed to save file: {e}")
    
    else:
        # Use intelligent naming with RAG system
        print(f"🧠 Using intelligent file naming with RAG analysis (naming_option: {naming_option})...")
        
        # First save with temp name for RAG analysis
        temp_path = os.path.join(upload_dir, f"temp_{filename}")
        with open(temp_path, 'wb') as f:
            f.write(file_content)
        
        try:
            # Generate intelligent filename
            new_filename = generate_intelligent_filename_optimized(
                rag_system=rag_system,
                original_filename=filename,
                naming_option=naming_option,
                user_title=user_title,
                user_client_name=user_client_name,
                counter=counter
            )
            
            # Create final file path
            final_path = os.path.join(upload_dir, new_filename)
            
            # Handle file conflicts
            if os.path.exists(final_path) and final_path != temp_path:
                base_name, ext = os.path.splitext(new_filename)
                conflict_counter = 1
                while os.path.exists(final_path):
                    final_path = os.path.join(upload_dir, f"{base_name}_{conflict_counter}{ext}")
                    conflict_counter += 1
                new_filename = os.path.basename(final_path)
                print(f"🔄 Resolved naming conflict: {new_filename}")
            
            # Move from temp to final location
            if final_path != temp_path:
                shutil.move(temp_path, final_path)
            
            print(f"✅ File saved with intelligent naming:")
            print(f"   Original: {filename}")
            print(f"   Final: {new_filename}")
            print(f"   Path: {final_path}")
            
            return final_path
            
        except Exception as e:
            print(f"❌ Error in intelligent file naming: {e}")
            # Fallback: use original temp file
            fallback_path = os.path.join(upload_dir, filename)
            if temp_path != fallback_path:
                shutil.move(temp_path, fallback_path)
            print(f"⚠️ Using fallback filename: {filename}")
            return fallback_path
        finally:
            # Clean up temp file if it still exists
            if os.path.exists(temp_path):
                try:
                    os.remove(temp_path)
                except:
                    pass

# ================================
# OPTIMIZED SINGLE-QUERY NAMING
# ================================

def generate_intelligent_filename_optimized(
    rag_system: Dict[str, Any],
    original_filename: str,
    naming_option: str,
    user_title: Optional[str] = None,
    user_client_name: Optional[str] = None,
    counter: Optional[int] = None
) -> str:
    """
    ⚡ ULTRA-OPTIMIZED: LLM returns the formatted filename directly.
    Single query with format instructions - no parsing/regex needed!
    """
    file_ext = os.path.splitext(original_filename)[1].lower()
    
    if naming_option == "keep_original":
        return original_filename
    
    # For intelligent naming options, use RAG with direct formatting
    if naming_option in ["add_timestamp", "add_client_name", "sequential_numbering"]:
        try:
            print(f"🚀 Generating {naming_option} filename with LLM direct formatting...")
            
            query_engine = rag_system.get("query_engine")
            if not query_engine:
                print("❌ No query engine found, using fallback naming")
                return generate_fallback_filename(original_filename, naming_option, user_title, user_client_name, counter)
            
            # Get current date for formatting
            current_date = datetime.now().strftime("%Y%m%d")
            
            # ⚡ SINGLE QUERY WITH DIRECT FORMAT INSTRUCTIONS
            if naming_option == "add_timestamp":
                format_query = f"""
Analyze this document and generate a filename in this EXACT format:
{current_date}_DOCUMENT-TYPE

Rules:
- Use today's date: {current_date}
- DOCUMENT-TYPE should be the document type (Contract, Agreement, Invoice, etc.) in PascalCase
- Remove all spaces, use uppercase for the document type
- Only respond with the filename, nothing else

Example format: {current_date}_SERVICE_AGREEMENT
Your response (filename only):"""

            elif naming_option == "add_client_name":
                format_query = f"""
Analyze this document and generate a filename in this EXACT format:
{current_date}_DOCUMENT-TYPE_SURNAME

Rules:
- Use today's date: {current_date}
- DOCUMENT-TYPE should be the document type (Contract, Agreement, Invoice, etc.) all UpperCase
- SURNAME is the LAST NAME ONLY of the main person/client mentioned (all UPPERCASE)
- Remove all spaces between words in document type
- Only respond with the filename, nothing else

Example format: {current_date}_SERVICE-CONTRACT_SMITH
Your response (filename only):"""

            elif naming_option == "sequential_numbering":
                seq_num = f"{counter:03d}" if counter else "001"
                format_query = f"""
Analyze this document and generate a filename in this EXACT format:
PERSON_DOCUMENT-TYPE_{seq_num}

Rules:
- PERSON is the main person/client name (FirstLast, no spaces)
- DOCUMENTTYPE should be the document type (Contract, Agreement, etc.) in Uppercase with spaces separated by -
- Use the sequence number: {seq_num}
- Remove all spaces
- Only respond with the filename, nothing else

Example format: SMITH_SERVICE-CONTRACT_{seq_num}
Your response (filename only):"""
            
            else:
                return generate_fallback_filename(original_filename, naming_option, user_title, user_client_name, counter)
            
            try:
                # ⚡ SINGLE API CALL - LLM returns formatted filename directly
                print("🤖 Requesting formatted filename from LLM...")
                response = query_engine.query(format_query).response.strip()
                
                # Clean up response (remove any extra text, keep just the filename)
                formatted_filename = clean_llm_filename_response(response)
                
                # Add file extension
                final_filename = f"{formatted_filename}{file_ext}"
                
                print(f"✅ LLM generated filename: {final_filename}")
                
                # Validate filename
                if is_valid_filename(formatted_filename):
                    return final_filename
                else:
                    print(f"⚠️ Invalid filename from LLM, using fallback")
                    return generate_fallback_filename(original_filename, naming_option, user_title, user_client_name, counter)
                
            except Exception as e:
                print(f"⚠️ LLM query failed: {e}, using fallback")
                return generate_fallback_filename(original_filename, naming_option, user_title, user_client_name, counter)
            
        except Exception as e:
            print(f"❌ Error in intelligent naming: {e}")
            print("⚠️ Falling back to user settings or original name")
            
    # Fallback to user-provided info or original name
    return generate_fallback_filename(original_filename, naming_option, user_title, user_client_name, counter)

def clean_llm_filename_response(response: str) -> str:
    """
    Clean LLM response to extract just the filename.
    Removes explanations, extra text, and ensures filename safety.
    """
    if not response:
        return ""
    
    # Split by newlines and take the first line (LLM might add explanations)
    lines = response.strip().split('\n')
    filename = lines[0].strip()
    
    # Remove common prefixes the LLM might add
    prefixes_to_remove = [
        "filename:", "response:", "answer:", "here is:", "the filename is:",
        "filename would be:", "suggested filename:", "my response:"
    ]
    
    filename_lower = filename.lower()
    for prefix in prefixes_to_remove:
        if filename_lower.startswith(prefix):
            filename = filename[len(prefix):].strip()
    
    # Remove quotes if present
    filename = filename.strip('"').strip("'")
    
    # Remove any file extensions (we add them back)
    if '.' in filename:
        filename = os.path.splitext(filename)[0]
    
    # Clean special characters (keep only alphanumeric, dash, underscore)
    filename = re.sub(r'[^\w\-_]', '', filename)
    
    # Limit length to reasonable size
    if len(filename) > 80:
        filename = filename[:80]
    
    return filename

def is_valid_filename(filename: str) -> bool:
    """
    Validate that a filename is safe and reasonable.
    """
    if not filename:
        return False
    
    # Must have reasonable length
    if len(filename) < 3 or len(filename) > 100:
        return False
    
    # Must contain at least one underscore (for our formats)
    if '_' not in filename:
        return False
    
    # Must not contain invalid characters
    if re.search(r'[<>:"/\\|?*\x00-\x1f]', filename):
        return False
    
    # Must not be only numbers/dates
    if filename.replace('_', '').isdigit():
        return False
    
    return True

def parse_combined_response(response: str) -> dict:
    """
    ⚡ DEPRECATED: This function is no longer needed with direct formatting.
    Kept for backwards compatibility.
    """
    info = {'date': None, 'person': None, 'type': None}
    
    try:
        lines = response.strip().split('\n')
        
        for line in lines:
            line = line.strip()
            if ':' in line:
                key, value = line.split(':', 1)
                key = key.strip().upper()
                value = value.strip()
                
                # Parse each component
                if key == 'DATE':
                    info['date'] = parse_date_from_response(value)
                elif key == 'PERSON':
                    info['person'] = parse_person_from_response(value)
                elif key == 'TYPE':
                    info['type'] = parse_type_from_response(value)
        
        print(f"📊 Parsed combined response:")
        print(f"   - Raw response: {response[:100]}...")
        print(f"   - Date parsed: {info['date']}")
        print(f"   - Person parsed: {info['person']}")
        print(f"   - Type parsed: {info['type']}")
        
        return info
        
    except Exception as e:
        print(f"❌ Error parsing combined response: {e}")
        print(f"❌ Raw response was: {response}")
        return {'date': None, 'person': None, 'type': None}

# ================================
# DEPRECATED PARSING FUNCTIONS
# (Kept for backwards compatibility only)
# ================================

def parse_date_from_response(response: str) -> Optional[str]:
    """Parse date from RAG response with multiple fallback patterns."""
    if not response or response.upper() in ['NO_DATE', 'NONE', 'UNKNOWN', 'N/A']:
        return None
    
    # Clean the response
    response = response.strip().upper()
    
    # Look for date patterns in the response
    date_patterns = [
        r'(\d{4})-(\d{1,2})-(\d{1,2})',  # YYYY-MM-DD
        r'(\d{4})/(\d{1,2})/(\d{1,2})',  # YYYY/MM/DD  
        r'(\d{1,2})/(\d{1,2})/(\d{4})',  # MM/DD/YYYY
        r'(\d{1,2})-(\d{1,2})-(\d{4})',  # MM-DD-YYYY
        r'(\d{4})(\d{2})(\d{2})',        # YYYYMMDD
    ]
    
    for pattern in date_patterns:
        match = re.search(pattern, response)
        if match:
            groups = match.groups()
            try:
                if len(groups) == 3:
                    # Determine format and convert to YYYYMMDD
                    if len(groups[0]) == 4:  # First group is year
                        year, month, day = int(groups[0]), int(groups[1]), int(groups[2])
                    else:  # First group is month/day
                        if int(groups[2]) > 31:  # Last group is year
                            month, day, year = int(groups[0]), int(groups[1]), int(groups[2])
                        else:
                            day, month, year = int(groups[0]), int(groups[1]), int(groups[2])
                    
                    # Validate date
                    if 1900 <= year <= 2100 and 1 <= month <= 12 and 1 <= day <= 31:
                        return f"{year:04d}{month:02d}{day:02d}"
            except (ValueError, IndexError):
                continue
    
    return None

def parse_person_from_response(response: str) -> Optional[str]:
    """Parse person name from RAG response."""
    if not response or response.upper() in ['NO_PERSON', 'NONE', 'UNKNOWN', 'N/A']:
        return None
    
    # Clean and validate the response
    cleaned = re.sub(r'[^\w\s-]', '', response.strip())
    cleaned = re.sub(r'\s+', ' ', cleaned).strip()
    
    # Filter out common non-name responses
    invalid_responses = [
        'i am sorry', 'sorry', 'the document', 'not found', 'not mentioned',
        'cannot find', 'no person', 'unknown', 'unclear', 'not specified',
        'the context', 'information', 'document date', 'no clear', 'main person'
    ]
    
    cleaned_lower = cleaned.lower()
    if any(invalid in cleaned_lower for invalid in invalid_responses):
        return None
    
    # Extract actual names (basic validation)
    if len(cleaned) >= 2 and len(cleaned) <= 50:
        # Remove common prefixes/suffixes that might appear
        words = cleaned.split()
        if len(words) >= 1:
            # Take first reasonable name-like part
            return ''.join(words[:2])  # Take first two words max, remove spaces
    
    return None

def parse_type_from_response(response: str) -> Optional[str]:
    """Parse document type from RAG response."""
    if not response or response.upper() in ['UNKNOWN', 'NONE', 'N/A']:
        return None
    
    # Clean the response
    cleaned = re.sub(r'[^\w\s]', ' ', response.strip())
    cleaned = re.sub(r'\s+', ' ', cleaned).strip()
    
    # Common document types to look for
    document_types = {
        'contract': 'Contract',
        'agreement': 'Agreement', 
        'invoice': 'Invoice',
        'receipt': 'Receipt',
        'letter': 'Letter',
        'power of attorney': 'PowerOfAttorney',
        'poa': 'PowerOfAttorney',
        'claim': 'Claim',
        'approval': 'Approval',
        'certificate': 'Certificate',
        'license': 'License',
        'permit': 'Permit',
        'application': 'Application',
        'form': 'Form',
        'statement': 'Statement',
        'report': 'Report',
        'mortgage': 'Mortgage',
        'deed': 'Deed',
        'title': 'Title',
        'insurance': 'Insurance',
        'security': 'Security',
        'instrument': 'Instrument',
        'partial': 'PartialApproval'
    }
    
    cleaned_lower = cleaned.lower()
    
    # Look for document type matches
    for key, value in document_types.items():
        if key in cleaned_lower:
            return value
    
    # If no match, try to extract first meaningful word
    words = cleaned.split()
    if words:
        first_word = words[0].strip()
        if len(first_word) >= 3 and first_word.isalpha():
            return first_word.capitalize()
    
    return None

def create_filename_from_extracted_info(
    info: Dict[str, Optional[str]], 
    original_filename: str, 
    naming_option: str, 
    counter: Optional[int] = None,
    user_title: Optional[str] = None,
    user_client_name: Optional[str] = None
) -> str:
    """
    ⚡ DEPRECATED: No longer used with direct LLM formatting.
    Create filename from extracted information with smart fallbacks.
    Kept for backwards compatibility only.
    """
    file_ext = os.path.splitext(original_filename)[1].lower()
    
    # Apply fallbacks for missing information
    date_part = info.get('date') or datetime.now().strftime("%Y%m%d")
    person_part = info.get('client_name') or info.get('person') or user_client_name or extract_name_from_filename(original_filename)
    type_part = info.get('document_type') or info.get('type') or user_title or extract_type_from_filename(original_filename)
    
    # Clean components
    if person_part:
        person_part = clean_component_for_filename(person_part)
    if type_part:
        type_part = clean_component_for_filename(type_part)
    
    print(f"🔧 Building filename with fallbacks:")
    print(f"   - Date: {date_part} (extracted: {info.get('date')})")
    print(f"   - Client: {person_part} (extracted: {info.get('client_name') or info.get('person')})")
    print(f"   - Type: {type_part} (extracted: {info.get('document_type') or info.get('type')})")
    
    # Build filename based on naming option
    if naming_option == "add_timestamp":
        # Format: YYYYMMDD_PersonName_DocumentType.ext
        parts = [date_part]
        if person_part:
            parts.append(person_part)
        if type_part:
            parts.append(type_part)
        
        if len(parts) >= 2:  # At least date + one other component
            final_filename = f"{'_'.join(parts)}{file_ext}"
        else:
            print("⚠️ Insufficient components for intelligent naming, using original")
            return original_filename
            
    elif naming_option == "sequential_numbering":
        # Format: PersonName_DocumentType_001.ext
        seq_num = f"{counter:03d}" if counter else "001"
        parts = []
        if person_part:
            parts.append(person_part)
        if type_part:
            parts.append(type_part)
        parts.append(seq_num)
        
        if len(parts) >= 2:  # At least one component + seq_num
            final_filename = f"{'_'.join(parts)}{file_ext}"
        else:
            print("⚠️ Insufficient components for intelligent naming, using original")
            return original_filename
    else:
        final_filename = original_filename
    
    # Ensure filename isn't too long
    if len(final_filename) > 100:
        final_filename = final_filename[:97] + file_ext
    
    return final_filename

def clean_component_for_filename(component: str) -> str:
    """Clean a component to be filename-safe."""
    if not component:
        return ""
    
    # Remove special characters and clean
    cleaned = re.sub(r'[^\w\s-]', '', component)
    cleaned = re.sub(r'\s+', '', cleaned)  # Remove all spaces
    cleaned = cleaned.strip()
    
    # Limit length
    return cleaned[:20] if cleaned else ""

def extract_name_from_filename(filename: str) -> Optional[str]:
    """Try to extract a name from the original filename."""
    base_name = os.path.splitext(filename)[0]
    
    # Common patterns in filenames
    name_patterns = [
        r'([A-Z][a-z]+ [A-Z][a-z]+)',  # FirstName LastName
        r'([A-Z][a-z]+_[A-Z][a-z]+)',  # FirstName_LastName
        r'_([A-Z][a-z]+ [A-Z][a-z]+)',  # _FirstName LastName
    ]
    
    for pattern in name_patterns:
        match = re.search(pattern, base_name)
        if match:
            name = match.group(1).replace('_', ' ').replace(' ', '')
            if len(name) >= 4:  # Reasonable name length
                return name
    
    return None

def extract_type_from_filename(filename: str) -> Optional[str]:
    """Try to extract document type from the original filename."""
    base_name = os.path.splitext(filename)[0].lower()
    
    # Common document type patterns
    type_patterns = {
        'contract': 'Contract',
        'agreement': 'Agreement',
        'power': 'PowerOfAttorney',
        'attorney': 'PowerOfAttorney',
        'poa': 'PowerOfAttorney',
        'invoice': 'Invoice',
        'claim': 'Claim',
        'approval': 'Approval',
        'partial': 'PartialApproval',
        'security': 'Security',
        'instrument': 'Instrument',
        'mortgage': 'Mortgage',
        'deed': 'Deed'
    }
    
    for keyword, doc_type in type_patterns.items():
        if keyword in base_name:
            return doc_type
    
    return None

def generate_fallback_filename(
    original_filename: str,
    naming_option: str,
    user_title: Optional[str] = None,
    user_client_name: Optional[str] = None,
    counter: Optional[int] = None
) -> str:
    """Enhanced fallback filename generation - surname at the end for add_client_name."""
    file_ext = os.path.splitext(original_filename)[1].lower()
    
    def clean_for_filename(text: str) -> str:
        if not text:
            return ""
        
        # Remove special characters
        cleaned = re.sub(r'[^\w\s-]', '', text)
        # Replace multiple spaces with single space
        cleaned = re.sub(r'\s+', ' ', cleaned).strip()
        
        # Filter out unknown words (case-insensitive)
        words_to_remove = ['page', 'unknown', 'document', 'file']
        words = cleaned.split()
        filtered_words = [w for w in words if w.lower() not in words_to_remove]
        
        if filtered_words:
            # Convert to uppercase and join with dash
            return '-'.join(filtered_words).upper()[:30]
        else:
            # If all words were filtered, return uppercase version
            return re.sub(r'\s+', '-', cleaned).upper()[:30]
    
    def extract_surname(text: str) -> str:
        """Extract only the surname (last word) from client name."""
        if not text:
            return ""
        # Split and get last word (surname)
        words = text.strip().split()
        return words[-1].upper() if words else ""
    
    # Try to use user settings if available
    if naming_option == "add_timestamp" and user_title:
        date_str = datetime.now().strftime("%Y%m%d")
        clean_title = clean_for_filename(user_title)
        return f"{date_str}_{clean_title}{file_ext}"
    
    elif naming_option == "add_client_name" and user_title and user_client_name:
        date_str = datetime.now().strftime("%Y%m%d")
        clean_title = clean_for_filename(user_title)
        surname = extract_surname(user_client_name)
        # Format: YYYYMMDD_DOCUMENTTYPE_SURNAME.ext
        return f"{date_str}_{clean_title}_{surname}{file_ext}"
    
    # Ultimate fallback: try to extract from filename, else use original
    elif naming_option in ["add_timestamp", "add_client_name"]:
        extracted_name = extract_name_from_filename(original_filename)
        extracted_type = extract_type_from_filename(original_filename)
        
        if naming_option == "add_timestamp" and extracted_type:
            date_str = datetime.now().strftime("%Y%m%d")
            return f"{date_str}_{clean_for_filename(extracted_type)}{file_ext}"
        
        elif naming_option == "add_client_name" and (extracted_name or extracted_type):
            date_str = datetime.now().strftime("%Y%m%d")
            parts = [date_str]
            
            if extracted_type:
                parts.append(clean_for_filename(extracted_type))
            
            if extracted_name:
                # Extract surname from the name
                surname = extract_surname(extracted_name)
                parts.append(surname)
            
            if len(parts) >= 2:
                return f"{'_'.join(parts)}{file_ext}"
    
    print(f"⚠️ No valid naming configuration, keeping original: {original_filename}")
    return original_filename

# ================================
# BACKWARD COMPATIBILITY
# ================================

# Keep the old function name for backward compatibility
def generate_intelligent_filename_simple(
    rag_system: Dict[str, Any],
    original_filename: str,
    naming_option: str,
    user_title: Optional[str] = None,
    user_client_name: Optional[str] = None,
    counter: Optional[int] = None
) -> str:
    """
    Backward compatibility wrapper for the optimized function.
    This ensures existing code still works while using the new optimized single-query approach.
    """
    print("⚡ Using optimized single-query intelligent naming...")
    return generate_intelligent_filename_optimized(
        rag_system=rag_system,
        original_filename=original_filename,
        naming_option=naming_option,
        user_title=user_title,
        user_client_name=user_client_name,
        counter=counter
    )

def validate_pdf_content(pdf_path: str) -> dict:
    """
    Validate that PDF has extractable text content.
    Returns validation info.
    """
    try:
        with fitz.open(pdf_path) as doc:
            total_text = ""
            pages_with_text = 0

            for page in doc:
                page_text = page.get_text().strip()
                total_text += page_text
                if len(page_text) > 10:  # Page has meaningful text
                    pages_with_text += 1

            validation_info = {
                "is_valid": len(total_text.strip()) > 20,
                "total_pages": len(doc),
                "pages_with_text": pages_with_text,
                "total_characters": len(total_text),
                "avg_chars_per_page": len(total_text) / max(len(doc), 1),
                "text_coverage": pages_with_text / max(len(doc), 1)
            }

            print(f"📊 PDF validation:")
            print(f"   Valid: {validation_info['is_valid']}")
            print(f"   Pages: {validation_info['total_pages']}")
            print(f"   Pages with text: {validation_info['pages_with_text']}")
            print(f"   Total characters: {validation_info['total_characters']}")

            return validation_info

    except Exception as e:
        print(f"❌ PDF validation failed: {e}")
        return {
            "is_valid": False,
            "error": str(e)
        }

def extract_text_from_pdf(pdf_path: str) -> List[Document]:
    """
    Extract text from PDF using direct extraction (OCR functions removed).
    Returns a list of Document objects.
    """
    print("✅ Document Type: Well-Structured (Direct Text Extraction)")
    documents = []
    
    with fitz.open(pdf_path) as doc:
        for i, page in enumerate(doc):
            # Extract raw text from each page
            text = page.get_text()

            # Preserve critical newlines for logical chunking
            text = "\n".join([line.strip() for line in text.splitlines() if line.strip()])

            # Create Document object with metadata
            documents.append(Document(
                text=text,
                metadata={
                    "file_name": os.path.basename(pdf_path),
                    "page_number": i + 1,
                    "total_pages": len(doc)
                }
            ))

    return documents

# ================================
# RATE LIMITING PROTECTION
# ================================

def is_rate_limit_safe() -> bool:
    """
    Check if it's safe to make RAG queries based on rate limiting.
    Simple implementation - can be enhanced with actual rate tracking.
    """
    # This is a simple check - you can enhance with actual rate tracking
    # For now, always return True but log the optimization
    print("🛡️ Rate limit check: Using single optimized query")
    return True

# ================================
# PERFORMANCE MONITORING
# ================================

def log_naming_performance(operation: str, duration: float, queries_used: int = 1):
    """
    Log performance metrics for intelligent naming operations.
    """
    print(f"📊 Naming Performance:")
    print(f"   - Operation: {operation}")
    print(f"   - Duration: {duration:.2f}s")
    print(f"   - Queries used: {queries_used}")
    print(f"   - Rate limit friendly: {queries_used == 1}")
